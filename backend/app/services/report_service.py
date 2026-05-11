"""Serviço de geração de relatórios PDF/DOCX (Fase 5 PRD v2).

Seis tipos de relatório, dois formatos cada:

- ``executive_summary`` — 1 página com score, probabilidade,
  forças/fraquezas, próxima ação, alertas.
- ``factor_deep_dive`` — detalhamento dos 12 fatores com origem.
- ``candidate_comparison`` — resultado Monte Carlo lado a lado.
- ``scenario_what_if`` — baseline vs alternativo de um cenário.
- ``compliance_audit`` — alertas LGPD/TSE abertos.
- ``dossier_export`` — dossiê completo de um candidato.

Stack:
- Jinja2 renderiza HTML (templates em ``app/templates/reports/``).
- WeasyPrint converte HTML→PDF. **Import lazy + try/except**: quando
  as libs de sistema (libpango/libcairo) não estão presentes (ex:
  CI de teste), o serviço marca ``PDF_AVAILABLE=False`` e o endpoint
  responde 503 para PDF. DOCX continua funcionando.
- python-docx renderiza um DOCX simples (sem 1:1 do HTML — gera
  estrutura nativa Word com headings + tabelas).

Branding: extraído de ``political_projects`` (header_logo_url,
footer_logo_url, candidate_photo_url + nome de urna/partido implícitos
via project.candidate_name/parties). Disclaimer fixo no rodapé.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

logger = logging.getLogger(__name__)

# Lazy import — em CI sem libpango, importar weasyprint quebra com OSError.
try:
    from weasyprint import HTML  # type: ignore

    PDF_AVAILABLE = True
except Exception as exc:  # noqa: BLE001
    HTML = None  # type: ignore
    PDF_AVAILABLE = False
    logger.warning("weasyprint_unavailable", extra={"reason": repr(exc)})

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt

    DOCX_AVAILABLE = True
except Exception as exc:  # noqa: BLE001
    Document = None  # type: ignore
    Pt = None  # type: ignore
    DOCX_AVAILABLE = False
    logger.warning("python_docx_unavailable", extra={"reason": repr(exc)})


_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates" / "reports"
_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


REPORT_TYPES = {
    "executive_summary": "executive_summary.html",
    "factor_deep_dive": "factor_deep_dive.html",
    "candidate_comparison": "candidate_comparison.html",
    "scenario_what_if": "scenario_what_if.html",
    "compliance_audit": "compliance_audit.html",
    "dossier_export": "dossier_export.html",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _confidence_class(level: str | None) -> str:
    return {"high": "high", "medium": "medium", "low": "low"}.get(
        (level or "").lower(), "medium"
    )


def _branding_block(project, *, disclaimer: str | None = None) -> dict[str, Any]:
    return {
        "header_logo_url": getattr(project, "header_logo_url", None),
        "footer_logo_url": getattr(project, "footer_logo_url", None),
        "candidate_photo_url": getattr(project, "candidate_photo_url", None),
        "footer_text": (
            f"{project.candidate_name} · "
            f"{', '.join(project.parties) if project.parties else ''} · "
            "Documento de uso interno"
        ),
        "disclaimer": disclaimer
        or (
            "Gerado por CampanhaProCenarios. Estimativas estatísticas, não "
            "predições. Documento de uso interno. Conteúdo gerado por IA quando "
            "indicado — verifique informações sensíveis antes de uso operacional."
        ),
    }


def _common_context(project, title: str, **extras) -> dict[str, Any]:
    return {
        "title": title,
        "project": project,
        "branding": _branding_block(project),
        "now_pt": datetime.now().strftime("%d/%m/%Y %H:%M"),
        **extras,
    }


def render_html(report_type: str, context: dict[str, Any]) -> str:
    if report_type not in REPORT_TYPES:
        raise ValueError(f"Tipo de relatório desconhecido: {report_type!r}")
    template = _env.get_template(REPORT_TYPES[report_type])
    return template.render(**context)


def html_to_pdf(html: str) -> bytes:
    """Converte HTML em PDF via WeasyPrint. Lança RuntimeError se indisponível."""
    if not PDF_AVAILABLE or HTML is None:
        raise RuntimeError(
            "WeasyPrint não disponível neste ambiente "
            "(faltam libs do sistema: libpango/libcairo)."
        )
    buf = io.BytesIO()
    HTML(string=html).write_pdf(buf)
    return buf.getvalue()


def html_to_docx_simple(title: str, sections: list[tuple[str, str]]) -> bytes:
    """Gera um DOCX simples com título + (heading, body) pairs.

    Não é uma renderização 1:1 do HTML — Word tem layout próprio. A função
    aceita pares ``(heading, body_text)`` e cria estrutura nativa.
    """
    if not DOCX_AVAILABLE or Document is None:
        raise RuntimeError("python-docx não disponível neste ambiente.")
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.size = Pt(20)
    for h, body in sections:
        if h:
            doc.add_heading(h, level=2)
        if body:
            doc.add_paragraph(body)
    doc.add_paragraph(
        "Gerado por CampanhaProCenarios. Estimativas estatísticas, não predições. "
        "Documento de uso interno."
    ).italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Builders — cada um monta o context específico do tipo
# ---------------------------------------------------------------------------


def build_executive_summary_context(
    project,
    *,
    latest_factors: dict[str, Any] | None = None,
    latest_election: dict[str, Any] | None = None,
    alerts: list[Any] | None = None,
    factor_catalog: list[Any] | None = None,
) -> dict[str, Any]:
    forces: list[dict[str, Any]] = []
    weaknesses: list[dict[str, Any]] = []
    score_atual = None
    coverage_percent = 0.0
    confidence_level = "medium"

    if latest_factors:
        factors = latest_factors.get("factors") or {}
        coverage_percent = latest_factors.get("coverage_percent", 0.0)
        # Map fator → label/hint usando o catálogo (passado de fora).
        catalog = {f.key: f for f in (factor_catalog or [])}
        sortable = [
            {
                "key": k,
                "value": v,
                "label": catalog[k].label if k in catalog else k,
                "hint": catalog[k].recommendation_hint if k in catalog else "",
            }
            for k, v in factors.items()
        ]
        sortable.sort(key=lambda x: x["value"], reverse=True)
        forces = sortable[:3]
        weaknesses = sortable[-3:][::-1]
        # score "rápido": média ponderada dos fatores presentes
        if sortable:
            score_atual = round(sum(s["value"] for s in sortable) / len(sortable), 1)

    win_probability = None
    if latest_election and latest_election.get("output_results"):
        # Mostra o candidato com maior win_probability.
        top = max(
            latest_election["output_results"],
            key=lambda r: r.get("win_probability", 0),
        )
        wp = top.get("win_probability", 0)
        if wp >= 0.95:
            win_probability = f"≥95% · {top.get('candidate_name')}"
        elif wp <= 0.05:
            win_probability = f"≤5% · {top.get('candidate_name')}"
        else:
            win_probability = f"{wp * 100:.1f}% · {top.get('candidate_name')}"

    next_action = weaknesses[0]["hint"] if weaknesses else None

    return _common_context(
        project,
        title="Resumo Executivo",
        confidence_level=confidence_level,
        confidence_class=_confidence_class(confidence_level),
        score_atual=score_atual,
        coverage_percent=round(coverage_percent, 1),
        win_probability=win_probability,
        forces=forces,
        weaknesses=weaknesses,
        next_action=next_action,
        alerts=alerts or [],
    )


def build_factor_deep_dive_context(
    project,
    *,
    cache,
    factor_catalog,
) -> dict[str, Any]:
    factors_dict = cache.factors or {} if cache else {}
    sources = cache.sources_used or {} if cache else {}
    rows = []
    for f in factor_catalog:
        rows.append(
            {
                "key": f.key,
                "label": f.label,
                "recommendation_hint": f.recommendation_hint,
                "value": factors_dict.get(f.key),
                "sources": sources.get(f.key, []),
            }
        )
    coverage_percent = cache.coverage_percent if cache else 0.0
    return _common_context(
        project,
        title="Deep Dive — 12 Fatores Eleitorais",
        confidence_level="medium",
        confidence_class="medium",
        coverage_percent=round(coverage_percent, 1),
        reference_date=cache.reference_date.strftime("%d/%m/%Y") if cache else "—",
        factors=rows,
        warnings=cache.warnings if cache else [],
    )


def build_candidate_comparison_context(project, *, election_result) -> dict[str, Any]:
    return _common_context(
        project,
        title="Comparação de Candidatos — Monte Carlo",
        confidence_level=election_result.confidence_level,
        confidence_class=_confidence_class(election_result.confidence_level),
        iterations=election_result.iterations,
        seed=election_result.seed,
        two_rounds=any(
            r.get("second_round_qualification_probability") is not None
            for r in (election_result.output_results or [])
        ),
        results=election_result.output_results or [],
    )


def build_scenario_what_if_context(
    project,
    *,
    scenario,
    factor_breakdown: list[dict[str, Any]],
    recommendations: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    return _common_context(
        project,
        title=f"Cenário What-If — {scenario.name}",
        confidence_level=getattr(scenario, "confidence_level", "medium"),
        confidence_class=_confidence_class(getattr(scenario, "confidence_level", "medium")),
        scenario=scenario,
        factor_breakdown=factor_breakdown,
        recommendations=recommendations or [],
    )


def build_compliance_audit_context(project, *, alerts) -> dict[str, Any]:
    return _common_context(
        project, title="Auditoria de Compliance", alerts=alerts
    )


def build_dossier_export_context(project, *, dossier) -> dict[str, Any]:
    return _common_context(
        project,
        title=f"Dossiê — {dossier.candidate_name}",
        confidence_level=dossier.confidence_level,
        confidence_class=_confidence_class(dossier.confidence_level),
        dossier=dossier,
    )


# ---------------------------------------------------------------------------
# Build DOCX a partir do mesmo contexto (versão simplificada)
# ---------------------------------------------------------------------------


def build_docx(report_type: str, context: dict[str, Any]) -> bytes:
    """Constrói DOCX a partir do mesmo contexto, sem template HTML.

    Lê os campos essenciais do contexto e gera estrutura nativa Word.
    Para os 6 tipos, uma seleção pragmática de campos por tipo.
    """
    title = context.get("title", "Relatório")
    project = context["project"]
    sections: list[tuple[str, str]] = [
        (
            "Identificação",
            f"{project.candidate_name} · {project.office} · {project.election_year} · "
            f"{project.municipality or ''} {project.state or ''}",
        ),
        ("Gerado em", context.get("now_pt", "")),
    ]

    if report_type == "executive_summary":
        sections.append(("Confiança global", context.get("confidence_level", "")))
        sections.append(("Score atual", str(context.get("score_atual") or "—")))
        sections.append(("Probabilidade de vitória", str(context.get("win_probability") or "—")))
        sections.append((
            "Cobertura de dados reais",
            f"{context.get('coverage_percent', 0)}%",
        ))
        if context.get("forces"):
            sections.append(
                (
                    "Top 3 forças",
                    "\n".join(f"- {f['label']}: {f['value']}/100" for f in context["forces"]),
                )
            )
        if context.get("weaknesses"):
            sections.append(
                (
                    "Top 3 fraquezas",
                    "\n".join(
                        f"- {f['label']}: {f['value']}/100 — {f['hint']}"
                        for f in context["weaknesses"]
                    ),
                )
            )

    elif report_type == "factor_deep_dive":
        sections.append(("Cobertura", f"{context.get('coverage_percent', 0)}%"))
        lines = []
        for f in context.get("factors", []):
            val = f["value"] if f["value"] is not None else "—"
            lines.append(f"- {f['label']}: {val} (origem: {', '.join(f['sources']) or 'manual'})")
        sections.append(("Fatores", "\n".join(lines)))

    elif report_type == "candidate_comparison":
        lines = []
        for r in context.get("results", []):
            wp = r.get("win_probability", 0)
            if wp >= 0.95:
                wp_s = "≥95%"
            elif wp <= 0.05:
                wp_s = "≤5%"
            else:
                wp_s = f"{wp * 100:.1f}%"
            ci = r.get("share_ci_95_first_round", [0, 0])
            lines.append(
                f"- {r['candidate_name']}: vitória {wp_s} · "
                f"share médio {r.get('mean_share_first_round', 0) * 100:.1f}% "
                f"(IC {ci[0] * 100:.1f}%–{ci[1] * 100:.1f}%)"
            )
        sections.append(("Candidatos", "\n".join(lines)))

    elif report_type == "scenario_what_if":
        s = context["scenario"]
        sections.append(
            (
                "Resultado",
                f"Baseline: {getattr(s, 'baseline_score', None)} · "
                f"Alternativo: {getattr(s, 'alternative_score', None)} · "
                f"Delta: {getattr(s, 'delta', None)}",
            )
        )

    elif report_type == "compliance_audit":
        alerts = context.get("alerts") or []
        sections.append(("Total de alertas abertos", str(len(alerts))))
        for a in alerts:
            sections.append(
                (
                    f"[{getattr(a, 'severity', '?')}] {getattr(a, 'alert_type', '')}",
                    getattr(a, "message", ""),
                )
            )

    elif report_type == "dossier_export":
        d = context["dossier"]
        sections.append(("Tipo", "Próprio" if d.candidate_type == "own" else "Adversário"))
        if d.biography:
            sections.append(("Biografia", d.biography))
        if d.ficha_limpa_status:
            sections.append(("Ficha Limpa", d.ficha_limpa_status))
        if d.strength_drivers:
            sections.append(("Drivers de apoio", "\n".join(f"- {x}" for x in d.strength_drivers)))
        if d.rejection_drivers:
            sections.append(
                ("Drivers de rejeição", "\n".join(f"- {x}" for x in d.rejection_drivers))
            )

    return html_to_docx_simple(title, sections)
